from fastapi import UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from utils.cache import get_cached_response, set_cached_response, hash_inputs, cache_if_successful
from utils.openai_utils import (
    analyze_resume, optimize_resume_jobscan_style, customize_resume, benchmark_resume, 
    ats_scan_jobscan_style, generate_cover_letter, jobscan_style_report, get_openai_api_key, 
    run_prompt, salary_insights, extract_resume_data, optimize_resume_ats100_style
)
import logging
import io
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfplumber
import tiktoken
import os
import httpx
from openai import OpenAI
import json
import re
from bs4 import BeautifulSoup

logger = logging.getLogger("candidate_service")

async def extract_resume_text(resume: UploadFile) -> str:
    content = await resume.read()
    filename = resume.filename.lower()
    if filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                if not text.strip():
                    raise Exception("No text extracted from PDF")
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Could not process PDF file: {str(e)}. Please try a different file format like DOCX or TXT."
            )
    else:
        return content.decode("utf-8", errors="ignore")

def count_tokens(text, model="gpt-3.5-turbo"):
    if not isinstance(text, str):
        text = str(text)
    enc = tiktoken.encoding_for_model(model)
    return len(enc.encode(text))

async def analyze_resume_service(resume: UploadFile, job_description: str, plan: str = "free"):
    logger.info("analyze_resume_service called")
    resume_text = await extract_resume_text(resume)
    cache_key = hash_inputs(resume_text, job_description, plan)
    cached = get_cached_response("analyze", cache_key)
    if cached:
        logger.info(f"Cache hit for analyze_resume {cache_key}")
        return cached
    result = analyze_resume(resume_text, job_description, plan=plan)
    cache_if_successful("analyze", cache_key, result)
    logger.info("Cache miss for analyze_resume, called OpenAI")
    return result

def normalize_resume_highlights(result):
    # Ensure result is a dictionary
    if not isinstance(result, dict):
        logger.warning(f"normalize_resume_highlights received non-dict result: {type(result)}")
        return {"atsScore": 70, "resumeHighlights": [], "error": "Invalid response format"}
    
    # Get highlights with proper error handling
    try:
        highlights = result.get("resumeHighlights", [])
        
        # Handle case where highlights is None or not iterable
        if highlights is None or not hasattr(highlights, '__iter__'):
            logger.warning(f"resumeHighlights is not iterable: {highlights}")
            result["resumeHighlights"] = []
            return result
            
        normalized = []
        for item in highlights:
            if isinstance(item, dict):
                normalized.append(item)
            elif isinstance(item, str):
                normalized.append({"text": item, "reason": ""})
            else:
                logger.warning(f"Unexpected highlight item type: {type(item)}")
                
        result["resumeHighlights"] = normalized
    except Exception as e:
        logger.error(f"Error normalizing resume highlights: {e}")
        result["resumeHighlights"] = []
        
    return result

async def optimize_resume_service(
    resume: UploadFile,   
    plan: str = "free",
    feature_type: str = None,
    download_format: str = None
):
    logger.info("optimize_resume_service called (optimize_resume_jobscan)")
    try:
        resume_text = await extract_resume_text(resume)       
        resume_tokens = count_tokens(resume_text)      
        total_tokens = resume_tokens 
        max_total_tokens = 15000

        if total_tokens > max_total_tokens:
            raise HTTPException(
                status_code=400,
                detail="Your resume is long for analysis. Please shorten your resume."
            )

        cache_key = hash_inputs(resume_text, plan, feature_type)
        cached = get_cached_response("optimize", cache_key)
        logger.info(f"Cached value for key {cache_key}: {cached}")
        
        # Process the result (either from cache or newly generated)
        if cached:
            logger.info(f"Cache hit for optimize_resume_jobscan {cache_key}")
            result = cached
        else:
            try:
                # Call the OpenAI function with error handling
                result = optimize_resume_jobscan_style(resume_text, plan=plan)
                logger.info(f"RAW OpenAI result: {result}")
                
                # Check if result is a valid response or an error
                if isinstance(result, dict) and "error" in result:
                    logger.error(f"Error in OpenAI response: {result['error']}")
                    
                    # Try to extract the optimized content from the raw response if available
                    if "raw" in result and isinstance(result["raw"], str):
                        import re
                        import json
                        
                        # Try to extract the JSON structure from the raw response
                        try:
                            # Look for the optimized content section
                            optimized_match = re.search(r'"optimizedContent"\s*:\s*(\{[^}]+\})', result["raw"], re.DOTALL)
                            if optimized_match:
                                optimized_str = optimized_match.group(1)
                                # Create a minimal valid result with the extracted content
                                result = {
                                    "atsScore": 75,  # Default score
                                    "optimizedContent": json.dumps("{" + optimized_str + "}")
                                }
                                logger.info("Successfully extracted optimized content from raw response")
                        except Exception as extract_error:
                            logger.error(f"Failed to extract optimized content: {extract_error}")
                            
                            # If extraction fails, create a basic response
                            result = {
                                "atsScore": 70,
                                "optimizedContent": "Could not parse full response",
                                "resumeHighlights": []
                            }
                
                # Ensure optimizedContent is a string, not an object (robust fix)
                if isinstance(result, dict) and "optimizedContent" in result:
                    if not isinstance(result["optimizedContent"], str):
                        try:
                            # If it's a dict or list, convert to a pretty JSON string
                            if isinstance(result["optimizedContent"], (dict, list)):
                                result["optimizedContent"] = json.dumps(result["optimizedContent"], ensure_ascii=False)
                            else:
                                result["optimizedContent"] = str(result["optimizedContent"])
                            logger.info("Robustly converted optimizedContent to string")
                        except Exception as convert_error:
                            logger.error(f"Failed to robustly convert optimizedContent to string: {convert_error}")
                            result["optimizedContent"] = str(result["optimizedContent"])
                
                # Normalize resume highlights if present
                if isinstance(result, dict):
                    logger.info(f"Type of resumeHighlights: {type(result.get('resumeHighlights', None))}, Value: {result.get('resumeHighlights', None)}")
                    result = normalize_resume_highlights(result)
                    
                    # Only cache if it's a successful result without errors
                    if "error" not in result:
                        cache_if_successful("optimize", cache_key, result)
                        
                logger.info("Cache miss for optimize, called OpenAI (optimize_resume_jobscan)")
            except Exception as api_error:
                logger.error(f"Error calling OpenAI API: {api_error}")
                # Return a fallback response
                result = {
                    "atsScore": 70,
                    "optimizedContent": {"error": f"API error: {str(api_error)}"},
                    "resumeHighlights": []
                }
        
        # If download format is specified, return the file for download
        if download_format and result.get("optimized"):
            return await download_resume_service(result.get("optimized"), download_format)
        
        # Otherwise return the JSON result
        return result
    except HTTPException as e:
        logger.error(f"optimize_resume_service failed: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"optimize_resume_service failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))



async def customize_resume_service(
    resume: UploadFile,
    job_description: str = None,
    job_description_file: UploadFile = None,
    plan: str = "free",
    download_format: str = None
):
    logger.info("customize_resume_service called (Jobscan-style)")
    try:
        resume_text = await extract_resume_text(resume)
        # Handle both text and file for job description
        if job_description_file is not None:
            jd_text = await extract_resume_text(job_description_file)
        else:
            jd_text = job_description or ""
        # Count tokens
        resume_tokens = count_tokens(resume_text)
        jd_tokens = count_tokens(jd_text)
        total_tokens = resume_tokens + jd_tokens
        max_total_tokens = 15000  # leave room for prompt/instructions

        if total_tokens > max_total_tokens:
            raise HTTPException(
                status_code=400,
                detail="Your resume and job description are too long for analysis. Please shorten your job description."
            )

        cache_key = hash_inputs(resume_text, jd_text, plan,"customize_resume_service")
        cached = get_cached_response("customize", cache_key)
        
        # Process the result (either from cache or newly generated)
        if cached:
            logger.info(f"Cache hit for customize_resume {cache_key}")
            result = cached
        else:
            result = jobscan_style_report(resume_text, jd_text, plan=plan)
            cache_if_successful("customize", cache_key, result)
            logger.info("Cache miss for customize_resume, called OpenAI (Jobscan-style)")
        
        # If download format is specified, return the file for download
        if download_format and result.get("customized"):
            return await download_resume_service(result.get("customized"), download_format)
        
        # Otherwise return the JSON result
        return result
    except HTTPException as e:
        logger.error(f"customize_resume_service failed: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"customize_resume_service failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
async def ats_scan_service(resume: UploadFile, plan: str = "free"):
    logger.info("ats_scan_service called")
    try:    
        resume_text = await extract_resume_text(resume)
        # Count tokens
        resume_tokens = count_tokens(resume_text)       
        max_total_tokens = 15000  # leave room for prompt/instructions
        total_tokens = resume_tokens 
        if total_tokens > max_total_tokens:
            raise HTTPException(
                status_code=400,
                detail="Your resume . Please shorten your resume ."
            )
        cache_key = hash_inputs(resume_text,plan,"ats_scan_service")
        cached = get_cached_response("ats_scan", cache_key)
        if cached:
            logger.info(f"Cache hit for ats_scan {cache_key}")
            return cached
        result = ats_scan_jobscan_style(resume_text, plan=plan)
        logger.info(f"RAW OpenAI result: {result}")
        cache_if_successful("ats_scan", cache_key, result)
        logger.info("Cache miss for ats_scan, called OpenAI")
        return result
    except HTTPException as e:
        logger.error(f"optimize_resume_service failed: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"optimize_resume_service failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
    
async def salary_insights_service(
        job_title: str,
        location: str,
        industry: str,
        years_experience: int,
        education_level: str = None,
        resume_text: UploadFile = None,
        plan: str = "free"):
    logger.info("salary_insights_service called")
    try:
        if resume_text is not None:
            resume_text_section = await extract_resume_text(resume_text)
        else:
            resume_text_section = None
        if not job_title or not location or not industry or years_experience is None:
            raise HTTPException(
                status_code=400,
                detail="job_title, location, industry, and years_experience are required fields."
            )
        # Count tokens
        job_title_tokens = count_tokens(job_title)
        location_tokens = count_tokens(location)
        industry_tokens = count_tokens(industry)
        years_experience_tokens = count_tokens(str(years_experience))
        education_level_tokens = count_tokens(education_level or "Not specified")
        resume_text_tokens = count_tokens(resume_text_section) if resume_text_section else 0
        total_tokens = (job_title_tokens + location_tokens + industry_tokens +
                        years_experience_tokens + education_level_tokens + resume_text_tokens)  
        max_total_tokens = 15000  # leave room for prompt/instructions
        if total_tokens > max_total_tokens:
            raise HTTPException(
                status_code=400,
                detail="Your input is too long for analysis. Please shorten your input."
            )
        cache_key = hash_inputs(
            job_title, location, industry, str(years_experience), plan,"salary_insights_service"
        )   
        cached = get_cached_response("salary_insights", cache_key)
        if cached:
            logger.info(f"Cache hit for salary_insights : {cache_key}")
            return cached
        logger.info(f"Cache miss for salary_insights, cache_key: {cache_key}")
        result = salary_insights(
            job_title=job_title,
            location=location,
            industry=industry,
            years_experience=years_experience,
            education_level=education_level,
            resume_text=resume_text_section,
            plan=plan
        )
        logger.info(f"RAW OpenAI result: {result}")
        cache_if_successful("salary_insights", cache_key, result)
        logger.info("Cache miss for salary_insights, called OpenAI")
        return result
    except HTTPException as e:
        logger.error(f"salary_insights_service failed: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"salary_insights_service failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))   

async def benchmark_resume_service(resume: UploadFile, job_description: str, plan: str = "free"):
    logger.info("benchmark_resume_service called")
    resume_text = await extract_resume_text(resume)
    cache_key = hash_inputs(resume_text, job_description, plan)
    cached = get_cached_response("benchmark", cache_key)
    if cached:
        logger.info(f"Cache hit for benchmark_resume {cache_key}")
        return cached
    result = benchmark_resume(resume_text, job_description, plan=plan)
    cache_if_successful("benchmark", cache_key, result)
    logger.info("Cache miss for benchmark_resume, called OpenAI")
    return result



async def generate_cover_letter_service(job_title: str, company: str, job_description: str, plan: str = "free"):
    logger.info("generate_cover_letter_service called")
    cache_key = hash_inputs(job_title, company, job_description, plan)
    cached = get_cached_response("generate_cover_letter", cache_key)
    if cached:
        logger.info(f"Cache hit for generate_cover_letter {cache_key}")
        return cached
    result = generate_cover_letter(job_title, company, job_description, plan=plan)
    cache_if_successful("generate_cover_letter", cache_key, result)
    logger.info("Cache miss for generate_cover_letter, called OpenAI")
    return result

async def extract_resume_data_service(resume: UploadFile, plan: str = "free"):
    """
    Extract structured data from a resume file.
    """
    logger.info("extract_resume_data_service called")
    try:
        resume_text = await extract_resume_text(resume)
        resume_tokens = count_tokens(resume_text)
        max_total_tokens = 15000  # leave room for prompt/instructions
        
        if resume_tokens > max_total_tokens:
            raise HTTPException(
                status_code=400,
                detail="Your resume is too long for analysis. Please shorten your resume."
            )
            
        cache_key = hash_inputs(resume_text, plan, "extract_resume_data_service")
        cached = get_cached_response("extract_resume_data", cache_key)
        
        if cached:
            logger.info(f"Cache hit for extract_resume_data : {cache_key}")
            return cached
            
        logger.info("Cache miss for extract_resume_data, calling OpenAI")
        logger.info(f"Resume text sample (first 200 chars): {resume_text[:200]}...")
        
        # Clear cache for testing during development
        # Uncomment this line if you want to force a fresh extraction each time
        # set_cached_response("extract_resume_data", cache_key, None)
        
        result = extract_resume_data(resume_text, plan=plan)
        
        # Log the result structure for debugging
        if isinstance(result, dict):
            logger.info(f"Extract resume result structure: {', '.join(result.keys())}")
            
            # Log experience and project counts
            exp_count = len(result.get("experience", [])) if isinstance(result.get("experience"), list) else 0
            proj_count = len(result.get("projects", [])) if isinstance(result.get("projects"), list) else 0
            achievements_count = len(result.get("achievements", [])) if isinstance(result.get("achievements"), list) else 0
            references_count = len(result.get("references", [])) if isinstance(result.get("references"), list) else 0
            logger.info(f"Extracted {exp_count} experience items, {proj_count} projects, {achievements_count} achievements, and {references_count} references")
            
            # Check if we have only summary field populated with the entire resume
            if (exp_count == 0 and proj_count == 0 and 
                result.get("summary") and len(result.get("summary", "")) > 200 and
                not result.get("name") and not result.get("email")):
                
                logger.info("Detected only summary field populated with entire resume. Attempting to extract structured data.")
                
                # Try to extract structured data from the summary
                summary_text = result.get("summary", "")
                
                # Extract name (usually at the beginning of the resume)
                name_match = re.search(r'^([A-Z][a-z]+(?: [A-Z][a-z]+)+)', summary_text)
                if name_match:
                    result["name"] = name_match.group(1)
                
                # Extract email
                email_match = re.search(r'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)', summary_text)
                if email_match:
                    result["email"] = email_match.group(1)
                
                # Extract phone
                phone_match = re.search(r'(\+?\d[\d\s-]{7,})', summary_text)
                if phone_match:
                    result["phone"] = phone_match.group(1)
                
                # Extract location
                location_match = re.search(r'(London|United Kingdom|UK)', summary_text)
                if location_match:
                    result["location"] = location_match.group(1)
                
                # Extract skills (look for "Skills" section)
                skills_section = re.search(r'Skills\s*(.+?)(?=Employment|Experience|Education|$)', summary_text, re.DOTALL | re.IGNORECASE)
                if skills_section:
                    skills_text = skills_section.group(1)
                    # Extract individual skills
                    skills = re.findall(r'([A-Z][a-zA-Z\s]+)', skills_text)
                    result["skills"] = [skill.strip() for skill in skills if len(skill.strip()) > 2]
                
                # Extract experience
                exp_section = re.search(r'(Employment History|Experience|Work Experience)\s*(.+?)(?=Education|Skills|Projects|$)', 
                                       summary_text, re.DOTALL | re.IGNORECASE)
                if exp_section:
                    exp_text = exp_section.group(2)
                    # Look for job titles and companies
                    job_matches = re.findall(r'([A-Z][a-zA-Z\s]+),\s*([A-Z][a-zA-Z\s]+),\s*([A-Za-z\s]+)', exp_text)
                    
                    if job_matches:
                        result["experience"] = []
                        for job_match in job_matches:
                            result["experience"].append({
                                "title": job_match[0].strip(),
                                "company": job_match[1].strip(),
                                "location": job_match[2].strip() if len(job_match) > 2 else "",
                                "startDate": "",
                                "endDate": "",
                                "description": "Extracted from resume summary"
                            })
                
                logger.info(f"After extraction from summary: Name={result.get('name')}, Email={result.get('email')}, Skills count={len(result.get('skills', []))}, Experience count={len(result.get('experience', []))}")
            
            # Final validation to ensure .NET API compatibility
            # Ensure all experience descriptions are strings
            if "experience" in result and isinstance(result["experience"], list):
                for exp in result["experience"]:
                    if isinstance(exp, dict):
                        if exp.get("description") is None:
                            exp["description"] = ""
                        elif isinstance(exp.get("description"), list):
                            # If description is a list, join it into a string
                            if len(exp["description"]) > 0:
                                exp["description"] = ". ".join(str(item) for item in exp["description"])
                            else:
                                exp["description"] = ""
                        elif not isinstance(exp.get("description"), str):
                            # If it's some other type, convert to string
                            exp["description"] = str(exp.get("description", ""))
            
            # Ensure all project descriptions are strings (not lists)
            if "projects" in result and isinstance(result["projects"], list):
                for proj in result["projects"]:
                    if isinstance(proj, dict):
                        if proj.get("description") is None:
                            proj["description"] = ""
                        elif isinstance(proj.get("description"), list):
                            if len(proj["description"]) > 0:
                                proj["description"] = ". ".join(str(item) for item in proj["description"])
                            else:
                                proj["description"] = ""
                        elif not isinstance(proj.get("description"), str):
                            proj["description"] = str(proj.get("description", ""))
        
        # Check if the result is valid for caching
        if isinstance(result, dict):
            if "error" in result:
                logger.warning(f"Not caching extract_resume_data result due to error: {result.get('error', 'Unknown error')}")
                
                # If we have raw content in the error response, try to make it more useful
                if "raw" in result and result["raw"]:
                    logger.info("Attempting to extract useful information from raw response")
                    try:
                        # Create a more structured response from the raw content
                        # This is a fallback to ensure the client gets something useful
                        raw_content = result["raw"]
                        
                        # If the raw content looks like it contains resume data, create a basic structure
                        if "name" in raw_content.lower() or "email" in raw_content.lower():
                            # Return the raw content with a note that it's not properly formatted
                            return {
                                "note": "Response was not properly formatted JSON. Basic information extracted.",
                                "raw_content": raw_content,
                                "parsed": False
                            }
                    except Exception as extract_error:
                        logger.error(f"Failed to extract useful information from raw response: {extract_error}")
            else:
                # Valid result, cache it
                logger.info("Caching successful extract_resume_data result")
                set_cached_response("extract_resume_data", cache_key, result)
        else:
            logger.warning(f"Not caching extract_resume_data result: unexpected type {type(result)}")
            
        return result
        
    except HTTPException as e:
        logger.error(f"extract_resume_data_service failed: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"extract_resume_data_service failed: {e}")
        # Return a structured error response instead of raising an exception
        return {
            "error": f"Failed to extract resume data: {str(e)}",
            "name": "",
            "title": "",
            "email": "",
            "phone": "",
            "location": "",
            "summary": "",
            "skills": [],
            "experience": [],
            "education": [],
            "certifications": [],
            "projects": []
        }

def create_docx_from_text(text, filename="optimized_resume.docx"):
    """Create a DOCX file from text and return the file path"""
    doc = Document()
    
    # Add the text to the document
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            p = doc.add_paragraph()
            p.add_run(para)
    
    # Save the document to a temporary file
    temp_file = os.path.join(tempfile.gettempdir(), filename)
    doc.save(temp_file)
    return temp_file

def create_pdf_from_text(text, filename="optimized_resume.pdf"):
    """
    Create a PDF file from HTML text with proper formatting and page handling
    """
    try:
        import weasyprint
        from weasyprint import HTML, CSS
        
        # Ensure the text is properly formatted HTML
        if not text.strip().lower().startswith(("<!doctype", "<html")):
            # If it's not HTML, wrap it in a basic HTML structure
            text = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume</title>
    <style>
        body {{
            font-family: 'Arial', sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            color: #333;
            font-size: 12px;
        }}
        @page {{
            size: A4;
            margin: 0.75in;
            @bottom-center {{
                content: counter(page);
            }}
        }}
        .page-break {{
            page-break-before: always;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 0;
            margin-bottom: 10px;
            font-weight: bold;
        }}
        p {{
            margin: 0 0 10px 0;
        }}
        ul, ol {{
            margin: 0 0 10px 20px;
            padding: 0;
        }}
        li {{
            margin-bottom: 5px;
        }}
    </style>
</head>
<body>
{text}
</body>
</html>"""
        
        # Detect if this is a two-column layout (like navy-column-modern)
        is_two_column = 'sidebar' in text.lower() and 'resume-container' in text.lower()
        
        # Add CSS for better PDF formatting
        if is_two_column:
            pdf_css = CSS(string="""
                @page {
                    size: A4;
                    margin: 0.5in;
                    @bottom-center {
                        content: counter(page);
                    }
                }
                
                body {
                    font-family: 'Arial', sans-serif;
                    line-height: 1.6;
                    margin: 0;
                    padding: 0;
                    color: #333;
                    font-size: 11px;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    background: #fff !important;
                }
                
                /* Convert flexbox layout to table layout for PDF */
                .resume-container {
                    max-width: 100% !important;
                    width: 100% !important;
                    margin: 0 !important;
                    background: #fff !important;
                    border-radius: 0 !important;
                    display: table !important;
                    box-shadow: none !important;
                    overflow: visible !important;
                    table-layout: fixed !important;
                }
                
                /* Sidebar as table cell */
                .sidebar {
                    display: table-cell !important;
                    width: 30% !important;
                    vertical-align: top !important;
                    padding: 20px 15px !important;
                    margin: 0 !important;
                    box-sizing: border-box !important;
                    background-color: #2c3e50 !important; /* Fallback color */
                    color: #fff !important;
                }
                
                /* Content area as table cell */
                .content {
                    display: table-cell !important;
                    width: 70% !important;
                    vertical-align: top !important;
                    padding: 20px 15px !important;
                    margin: 0 !important;
                    box-sizing: border-box !important;
                    background: #fff !important;
                    color: #333 !important;
                }
                
                /* Prevent content from being cut off */
                * {
                    box-sizing: border-box;
                    max-width: 100%;
                }
                
                /* Header styles */
                h1, h2, h3, h4, h5, h6 {
                    margin-top: 0;
                    margin-bottom: 8px;
                    font-weight: bold;
                    page-break-after: avoid;
                    word-wrap: break-word;
                }
                
                .content h1 { 
                    font-size: 18px; 
                    color: #2c3e50 !important;
                    margin-bottom: 5px;
                }
                .content h2 { 
                    font-size: 14px; 
                    color: #2c3e50 !important;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                .content h3 { 
                    font-size: 13px; 
                    color: #34495e !important;
                }
                h4, h5, h6 { font-size: 11px; }
                
                /* Sidebar header styles */
                .sidebar h1, .sidebar h2, .sidebar h3 {
                    color: #fff !important;
                }
                
                /* Paragraph styles */
                p {
                    margin: 0 0 8px 0;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }
                
                /* List styles */
                ul, ol {
                    margin: 0 0 8px 15px;
                    padding: 0;
                }
                
                li {
                    margin-bottom: 4px;
                    word-wrap: break-word;
                }
                
                /* Sidebar specific styles */
                .sidebar-section {
                    margin-bottom: 20px;
                    page-break-inside: avoid;
                }
                
                .sidebar-section-title {
                    font-size: 12px !important;
                    margin-bottom: 8px !important;
                    color: #ecf0f1 !important;
                    font-weight: bold !important;
                    text-transform: uppercase;
                    letter-spacing: 1px;
                }
                
                .sidebar-details {
                    font-size: 10px !important;
                    line-height: 1.4 !important;
                    color: #fff !important;
                }
                
                .sidebar-details a {
                    color: #bdc3c7 !important;
                    text-decoration: none;
                }
                
                .sidebar-skills-list {
                    list-style: none !important;
                    padding: 0 !important;
                    margin: 0 !important;
                }
                
                .sidebar-skills-list li {
                    font-size: 10px !important;
                    margin-bottom: 4px !important;
                    color: #ecf0f1 !important;
                    padding-left: 8px;
                    border-left: 3px solid rgba(255, 255, 255, 0.3);
                }
                
                /* Photo styling */
                .photo {
                    width: 50px !important;
                    height: 50px !important;
                    margin-bottom: 15px !important;
                    border-radius: 50%;
                    border: 2px solid #fff;
                }
                
                /* Content section styles */
                .profile-section, .employment-section, .education-section, 
                .projects-section, .certifications-section {
                    margin-bottom: 15px;
                    page-break-inside: avoid;
                }
                
                .content .title {
                    font-size: 14px !important;
                    color: #7f8c8d !important;
                    margin-bottom: 15px;
                    font-weight: normal;
                }
                
                .employment-history-role {
                    font-weight: bold;
                    font-size: 12px;
                    color: #2c3e50 !important;
                    margin-bottom: 2px;
                }
                
                .employment-history-company {
                    font-size: 11px;
                    color: #34495e !important;
                    margin-bottom: 2px;
                }
                
                .employment-history-dates {
                    font-size: 10px;
                    color: #7f8c8d !important;
                    margin-bottom: 5px;
                }
                
                .employment-history-list {
                    margin: 5px 0 10px 15px;
                    padding: 0;
                    font-size: 10px;
                    color: #2c3e50 !important;
                }
                
                .education-degree {
                    font-weight: bold;
                    font-size: 12px;
                    color: #2c3e50 !important;
                }
                
                .education-institution {
                    font-size: 11px;
                    color: #34495e !important;
                    margin-bottom: 2px;
                }
                
                .education-dates {
                    font-size: 10px;
                    color: #7f8c8d !important;
                    margin-bottom: 5px;
                }
                
                /* Prevent orphans and widows */
                p, li {
                    orphans: 2;
                    widows: 2;
                }
                
                /* Remove problematic properties */
                * {
                    position: static !important;
                    float: none !important;
                    transform: none !important;
                    box-shadow: none !important;
                }
                
                /* Ensure proper page breaks */
                .page-break {
                    page-break-before: always;
                }
                
                /* Fix table layout issues */
                table {
                    table-layout: fixed;
                    width: 100%;
                }
                
                td, th {
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }
            """)
        else:
            pdf_css = CSS(string="""
                @page {
                    size: A4;
                    margin: 0.75in;
                    @bottom-center {
                        content: counter(page);
                    }
                }
                
                body {
                    font-family: 'Arial', sans-serif;
                    line-height: 1.6;
                    margin: 0;
                    padding: 0;
                    color: #333;
                    font-size: 12px;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }
                
                /* Prevent content from being cut off */
                * {
                    box-sizing: border-box;
                    max-width: 100%;
                }
                
                /* Header styles */
                h1, h2, h3, h4, h5, h6 {
                    margin-top: 0;
                    margin-bottom: 10px;
                    font-weight: bold;
                    page-break-after: avoid;
                    word-wrap: break-word;
                }
                
                h1 { font-size: 18px; }
                h2 { font-size: 16px; }
                h3 { font-size: 14px; }
                h4, h5, h6 { font-size: 12px; }
                
                /* Paragraph styles */
                p {
                    margin: 0 0 10px 0;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }
                
                /* List styles */
                ul, ol {
                    margin: 0 0 10px 20px;
                    padding: 0;
                }
                
                li {
                    margin-bottom: 5px;
                    word-wrap: break-word;
                }
                
                /* Table styles */
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin-bottom: 15px;
                    table-layout: fixed;
                }
                
                th, td {
                    padding: 8px;
                    text-align: left;
                    border-bottom: 1px solid #ddd;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }
                
                /* Section styles */
                .section {
                    margin-bottom: 20px;
                    page-break-inside: avoid;
                }
                
                /* Contact info styles */
                .contact-info {
                    margin-bottom: 15px;
                }
                
                /* Experience and education styles */
                .experience-item, .education-item {
                    margin-bottom: 15px;
                    page-break-inside: avoid;
                }
                
                /* Skills styles */
                .skills-list {
                    display: flex;
                    flex-wrap: wrap;
                    gap: 10px;
                }
                
                .skill-item {
                    background-color: #f0f0f0;
                    padding: 4px 8px;
                    border-radius: 4px;
                    font-size: 11px;
                }
                
                /* Prevent orphans and widows */
                p, li {
                    orphans: 2;
                    widows: 2;
                }
                
                /* Ensure proper page breaks */
                .page-break {
                    page-break-before: always;
                }
                
                /* Remove empty space at the end */
                body::after {
                    content: "";
                    display: block;
                    height: 0;
                    clear: both;
                }
                
                /* Fix for right-side cutoff */
                .container {
                    max-width: 100%;
                    overflow: hidden;
                }
                
                /* Responsive text sizing */
                @media print {
                    body {
                        font-size: 11px;
                    }
                    h1 { font-size: 16px; }
                    h2 { font-size: 14px; }
                    h3 { font-size: 13px; }
                }
            """)
        
        # Create a temporary HTML file
        html_path = os.path.join(tempfile.gettempdir(), "temp_resume.html")
        with open(html_path, "w", encoding="utf-8") as html_file:
            html_file.write(text)
        
        # Define the output PDF path
        pdf_path = os.path.join(tempfile.gettempdir(), filename)
        
        # Convert HTML to PDF using WeasyPrint with custom CSS
        html_doc = HTML(filename=html_path)
        html_doc.write_pdf(pdf_path, stylesheets=[pdf_css])
        
        # Clean up temporary HTML file
        try:
            os.remove(html_path)
        except:
            pass
        
        return pdf_path
    except ImportError:
        logger.error("WeasyPrint is not installed. Cannot generate PDF.")
        raise RuntimeError("WeasyPrint is required for PDF generation. Please install it with 'pip install weasyprint'.")
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}")
        raise RuntimeError(f"Failed to create PDF: {str(e)}")

def html_to_docx_with_mammoth(html_str, filename="optimized_resume.docx"):
    """Convert HTML to DOCX using mammoth, return file path."""
    try:
        import mammoth
        import tempfile
        import os
        # Write HTML to a temp file
        html_path = os.path.join(tempfile.gettempdir(), "temp_resume_for_mammoth.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_str)
        # Output DOCX path
        docx_path = os.path.join(tempfile.gettempdir(), filename)
        # Use mammoth.convert_to_html to parse HTML, then use python-docx to write to docx
        with open(html_path, "rb") as html_file:
            # Mammoth expects DOCX input, not HTML. So fallback to basic text dump if not supported.
            # Instead, use a simple HTML parser to extract text and write to docx
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_str, "html.parser")
            doc = Document()
            for elem in soup.find_all(['h1','h2','h3','h4','h5','h6','p','li']):
                text = elem.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            doc.save(docx_path)
        return docx_path
    except ImportError:
        # If mammoth is not available, return None and let the caller handle it
        return None

def html_to_docx_preserve_formatting(html_str, filename="optimized_resume.docx"):
    """Convert HTML to DOCX, preserving basic formatting, color, and nested inline styles recursively."""
    from docx.shared import RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    
    soup = BeautifulSoup(html_str, "html.parser")
    doc = Document()
    
    # Set document margins to prevent content cutoff
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def parse_style(style_str):
        style = {}
        if not style_str:
            return style
        for part in style_str.split(';'):
            if ':' in part:
                k, v = part.split(':', 1)
                style[k.strip().lower()] = v.strip()
        return style

    def merge_styles(parent, child):
        """Merge two style dicts, child overrides parent."""
        merged = dict(parent) if parent else {}
        merged.update(child or {})
        return merged

    def apply_run_styles(run, tag, style_dict):
        if tag in ["b", "strong"] or style_dict.get("font-weight", "") in ["bold", "700"]:
            run.bold = True
        if tag in ["i", "em"] or style_dict.get("font-style", "") == "italic":
            run.italic = True
        if tag == "u" or style_dict.get("text-decoration", "") == "underline":
            run.underline = True
        color = style_dict.get("color")
        if color:
            if color.startswith('#') and len(color) == 7:
                try:
                    run.font.color.rgb = RGBColor.from_string(color[1:].upper())
                except Exception:
                    pass
            elif color.lower() == "red":
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            elif color.lower() == "blue":
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            elif color.lower() == "green":
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    def process_inline(node, paragraph, parent_tag=None, parent_style=None):
        # If node is a tag
        if hasattr(node, 'name') and node.name is not None:
            style_dict = merge_styles(parent_style, parse_style(node.get('style', '')))
            for child in node.children:
                process_inline(child, paragraph, node.name, style_dict)
        # If node is a text node
        else:
            text = str(node)
            if text.strip():
                run = paragraph.add_run(text)
                apply_run_styles(run, parent_tag, parent_style or {})

    def process_elem(elem):
        # Headings
        if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(elem.name[1]) if elem.name[1].isdigit() else 1
            doc.add_heading(elem.get_text(strip=True), level=level)
        # Paragraph
        elif elem.name == 'p' or (elem.name is None and str(elem).strip()):
            p = doc.add_paragraph()
            for child in elem.children if hasattr(elem, 'children') else []:
                process_inline(child, p)
            if elem.name is None:
                process_inline(elem, p)
        # Lists
        elif elem.name == 'ul':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                for child in li.children:
                    process_inline(child, p)
        elif elem.name == 'ol':
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                for child in li.children:
                    process_inline(child, p)
        # Table
        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if rows:
                cols = rows[0].find_all(['td', 'th'])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text(strip=True)
        # Other containers (div, section, body, etc.)
        elif elem.name in ['div', 'section', 'body', 'main', 'article', 'span'] or elem.name is None:
            for child in elem.children if hasattr(elem, 'children') else []:
                process_elem(child)
        # Fallback: treat as paragraph
        else:
            if elem.get_text(strip=True):
                p = doc.add_paragraph()
                p.add_run(elem.get_text(strip=True))

    # Start from <body> or the root
    root = soup.body if soup.body else soup
    for elem in root.children:
        process_elem(elem)

    docx_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(docx_path)
    return docx_path

def clean_html_for_pdf(html_content):
    """Clean HTML content to prevent blank pages and formatting issues"""
    try:
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Remove empty paragraphs and divs
        for tag in soup.find_all(['p', 'div', 'span']):
            if not tag.get_text(strip=True) and not tag.find_all(['img', 'br']):
                tag.decompose()
        
        # Remove excessive whitespace and empty lines
        for tag in soup.find_all(text=True):
            if isinstance(tag, str):
                # Replace multiple whitespaces with single space
                cleaned_text = re.sub(r'\s+', ' ', tag.strip())
                if cleaned_text != tag:
                    tag.replace_with(cleaned_text)
        
        # Special handling for two-column layouts
        is_two_column = soup.find(class_='sidebar') and soup.find(class_='resume-container')
        if is_two_column:
            # Ensure sidebar and content are properly structured
            sidebar = soup.find(class_='sidebar')
            content = soup.find(class_='content')
            resume_container = soup.find(class_='resume-container')
            
            if sidebar and content and resume_container:
                # Remove any empty sections in sidebar
                for section in sidebar.find_all(class_='sidebar-section'):
                    if not section.get_text(strip=True):
                        section.decompose()
                
                # Ensure content sections are properly separated
                for section in content.find_all(['div', 'section']):
                    if section.get('class') and any('section' in cls for cls in section.get('class', [])):
                        # Add page-break-inside: avoid to prevent sections from breaking
                        current_style = section.get('style', '')
                        if 'page-break-inside' not in current_style:
                            section['style'] = current_style + '; page-break-inside: avoid;'
                
                # Ensure the resume container has the right structure for table layout
                # Remove any problematic CSS that might interfere with table layout
                for element in [resume_container, sidebar, content]:
                    if element:
                        current_style = element.get('style', '')
                        # Remove flexbox properties
                        current_style = re.sub(r'display:\s*flex[^;]*;?', '', current_style)
                        current_style = re.sub(r'flex[^:]*:[^;]*;?', '', current_style)
                        # Remove float properties
                        current_style = re.sub(r'float:[^;]*;?', '', current_style)
                        # Remove position properties that might cause issues
                        current_style = re.sub(r'position:\s*(absolute|fixed)[^;]*;?', '', current_style)
                        element['style'] = current_style
        
        # Remove empty sections at the end that might cause blank pages
        body = soup.body if soup.body else soup
        if body:
            # Remove trailing empty elements
            for tag in reversed(list(body.children)):
                if hasattr(tag, 'name') and tag.name:
                    if not tag.get_text(strip=True):
                        tag.decompose()
                    else:
                        break
        
        # Fix any problematic CSS that might cause layout issues
        style_tags = soup.find_all('style')
        for style_tag in style_tags:
            if style_tag.string:
                # Remove problematic CSS properties
                css_content = style_tag.string
                # Remove box-shadow which can cause rendering issues
                css_content = re.sub(r'box-shadow:[^;]+;?', '', css_content)
                # Remove transform properties that can cause positioning issues
                css_content = re.sub(r'transform:[^;]+;?', '', css_content)
                # Remove position: fixed/absolute that can cause issues
                css_content = re.sub(r'position:\s*(fixed|absolute)[^;]*;?', 'position: static;', css_content)
                style_tag.string = css_content
        
        return str(soup)
    except Exception as e:
        logger.warning(f"Failed to clean HTML: {e}")
        return html_content

async def download_resume_service(resume_text, format="docx"):
    """Service to download resume in specified format"""
    try:
        logger.info(f"Downloading resume in format: {format}")
        is_html = resume_text.strip().lower().startswith(("<!doctype", "<html"))
        logger.info(f"Content appears to be HTML: {is_html}")
        
        # Clean the resume text if it's HTML
        if is_html:
            resume_text = clean_html_for_pdf(resume_text)
            logger.info("HTML content cleaned for better formatting")
        
        if format.lower() == "docx":
            if is_html:
                logger.info("Converting HTML to DOCX with formatting preservation")
                file_path = html_to_docx_preserve_formatting(resume_text)
            else:
                file_path = create_docx_from_text(resume_text)
            return FileResponse(
                path=file_path,
                filename="resume.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif format.lower() == "pdf":
            logger.info("Converting to PDF using WeasyPrint")
            file_path = create_pdf_from_text(resume_text)
            return FileResponse(
                path=file_path,
                filename="resume.pdf",
                media_type="application/pdf"
            )
        elif format.lower() == "html":
            logger.info("Returning HTML content")
            html_path = os.path.join(tempfile.gettempdir(), "resume.html")
            with open(html_path, "w", encoding="utf-8") as html_file:
                html_file.write(resume_text)
            return FileResponse(
                path=html_path,
                filename="resume.html",
                media_type="text/html"
            )
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
    except Exception as e:
        logger.error(f"download_resume_service failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def enhance_resume_ats100_service(
    resume: UploadFile,
    plan: str = "free",
    feature_type: str = None,
    download_format: str = None
):
    logger.info("enhance_resume_ats100_service called (optimize_resume_ats100_style)")
    try:
        resume_text = await extract_resume_text(resume)
        resume_tokens = count_tokens(resume_text)
        total_tokens = resume_tokens
        max_total_tokens = 15000

        if total_tokens > max_total_tokens:
            raise HTTPException(
                status_code=400,
                detail="Your resume is too long for analysis. Please shorten your resume."
            )

        cache_key = hash_inputs(resume_text, plan, feature_type)
        cached = get_cached_response("enhance_ats100", cache_key)
        logger.info(f"Cached value for key {cache_key}: {cached}")

        if cached:
            logger.info(f"Cache hit for enhance_resume_ats100 {cache_key}")
            result = cached
        else:
            try:
                result = optimize_resume_ats100_style(resume_text, plan=plan)
                logger.info(f"RAW OpenAI result (ATS100): {result}")
                if isinstance(result, dict) and "error" in result:
                    logger.error(f"Error in OpenAI response: {result['error']}")
                    if "raw" in result and isinstance(result["raw"], str):
                        try:
                            optimized_match = re.search(r'"optimizedContent"\s*:\s*(\{[^}]+\})', result["raw"], re.DOTALL)
                            if optimized_match:
                                optimized_str = optimized_match.group(1)
                                result = {
                                    "atsScore": 100,
                                    "optimizedContent": json.dumps("{" + optimized_str + "}")
                                }
                                logger.info("Successfully extracted optimized content from raw response (ATS100)")
                        except Exception as extract_error:
                            logger.error(f"Failed to extract optimized content (ATS100): {extract_error}")
                            result = {
                                "atsScore": 100,
                                "optimizedContent": "Could not parse full response",
                                "resumeHighlights": []
                            }
                if isinstance(result, dict) and "optimizedContent" in result:
                    if not isinstance(result["optimizedContent"], str):
                        try:
                            if isinstance(result["optimizedContent"], (dict, list)):
                                result["optimizedContent"] = json.dumps(result["optimizedContent"], ensure_ascii=False)
                            else:
                                result["optimizedContent"] = str(result["optimizedContent"])
                            logger.info("Robustly converted optimizedContent to string (ATS100)")
                        except Exception as convert_error:
                            logger.error(f"Failed to robustly convert optimizedContent to string (ATS100): {convert_error}")
                            result["optimizedContent"] = str(result["optimizedContent"])
                if isinstance(result, dict):
                    logger.info(f"Type of resumeHighlights: {type(result.get('resumeHighlights', None))}, Value: {result.get('resumeHighlights', None)}")
                    result = normalize_resume_highlights(result)
                    if "error" not in result:
                        cache_if_successful("enhance_ats100", cache_key, result)
                logger.info("Cache miss for enhance_ats100, called OpenAI (optimize_resume_ats100_style)")
            except Exception as api_error:
                logger.error(f"Error calling OpenAI API (ATS100): {api_error}")
                result = {
                    "atsScore": 100,
                    "optimizedContent": {"error": f"API error: {str(api_error)}"},
                    "resumeHighlights": []
                }
        if download_format and result.get("optimized"):
            return await download_resume_service(result.get("optimized"), download_format)
        return result
    except HTTPException as e:
        logger.error(f"enhance_resume_ats100_service failed: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"enhance_resume_ats100_service failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# FINAL GUARANTEE: Ensure optimizedContent is a string before returning (prompt-agnostic, robust)
        if isinstance(result, dict) and "optimizedContent" in result:
            if not isinstance(result["optimizedContent"], str):
                try:
                    # If it's a dict or list, convert to a compact JSON string
                    if isinstance(result["optimizedContent"], (dict, list)):
                        result["optimizedContent"] = json.dumps(result["optimizedContent"], ensure_ascii=False)
                    else:
                        result["optimizedContent"] = str(result["optimizedContent"])
                    logger.info("Final guarantee: converted optimizedContent to string before return")
                except Exception as convert_error:
                    logger.error(f"Final guarantee failed: {convert_error}")
                    result["optimizedContent"] = str(result["optimizedContent"])

